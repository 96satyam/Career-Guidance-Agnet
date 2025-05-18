import os
from fpdf import FPDF
from docx import Document

class ResumeFileGenerator:
    def __init__(self, output_dir="outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_pdf(self, resume_data, filename="resume.pdf"):
        """
        Generates a PDF resume using fpdf.
        resume_data: dict with keys like name, education, experience, skills
        """
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt=resume_data.get("name", "Name"), ln=True, align="C")

        pdf.set_font("Arial", 'B', 14)
        pdf.cell(200, 10, txt="Education", ln=True)
        pdf.set_font("Arial", '', 12)
        for edu in resume_data.get("education", []):
            pdf.cell(200, 8, txt=edu, ln=True)

        pdf.set_font("Arial", 'B', 14)
        pdf.cell(200, 10, txt="Experience", ln=True)
        pdf.set_font("Arial", '', 12)
        for exp in resume_data.get("experience", []):
            pdf.cell(200, 8, txt=exp, ln=True)

        pdf.set_font("Arial", 'B', 14)
        pdf.cell(200, 10, txt="Skills", ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 8, txt=", ".join(resume_data.get("skills", [])))

        filepath = os.path.join(self.output_dir, filename)
        pdf.output(filepath)
        return filepath

    def generate_docx(self, resume_data, filename="resume.docx"):
        """
        Generates a DOCX resume using python-docx.
        resume_data: dict with keys like name, education, experience, skills
        """
        doc = Document()
        doc.add_heading(resume_data.get("name", "Name"), 0)

        doc.add_heading("Education", level=1)
        for edu in resume_data.get("education", []):
            doc.add_paragraph(edu)

        doc.add_heading("Experience", level=1)
        for exp in resume_data.get("experience", []):
            doc.add_paragraph(exp)

        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(resume_data.get("skills", [])))

        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath



